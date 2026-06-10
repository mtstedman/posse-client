#!/usr/bin/env python3
"""File format converter — converts between common file formats.

Usage:
    python tools/convert_file.py <input_file> <output_file> [options]

Supported conversions:
    PDF  -> TXT, JSON, images (PNG per page)
    DOCX -> TXT, JSON
    XLSX -> CSV, JSON
    CSV  -> JSON, XLSX
    JSON -> CSV, XLSX

Options:
    --pages 1-5     Page range (PDF conversions)
    --sheet Name    Sheet name (XLSX source)
    --dpi 150       DPI for PDF->image conversion (default: 150)
"""

import argparse
import csv
import json
import sys
from pathlib import Path


def pdf_to_text(src, dst, pages=None):
    import pdfplumber

    parts = []
    with pdfplumber.open(src) as pdf:
        page_range = _resolve_pages(pages, len(pdf.pages))
        for i in page_range:
            text = pdf.pages[i].extract_text()
            if text:
                parts.append(f"--- Page {i + 1} ---\n{text}")
    dst.write_text("\n\n".join(parts), encoding="utf-8")


def pdf_to_json(src, dst, pages=None):
    import pdfplumber

    result = {"pages": [], "tables": []}
    with pdfplumber.open(src) as pdf:
        page_range = _resolve_pages(pages, len(pdf.pages))
        for i in page_range:
            page = pdf.pages[i]
            text = page.extract_text()
            if text:
                result["pages"].append({"page": i + 1, "text": text})
            for ti, table in enumerate(page.extract_tables()):
                result["tables"].append({"page": i + 1, "index": ti, "rows": table})
    dst.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")


def pdf_to_images(src, dst_dir, pages=None, dpi=150):
    import fitz  # pymupdf

    dst_dir = Path(dst_dir)
    dst_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(str(src))
    page_range = _resolve_pages(pages, len(doc))
    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)
    created = []
    for i in page_range:
        pix = doc[i].get_pixmap(matrix=mat)
        out_path = dst_dir / f"page_{i + 1:03d}.png"
        pix.save(str(out_path))
        created.append(str(out_path))
    doc.close()
    print(f"Created {len(created)} page image(s) in {dst_dir}")
    for p in created:
        print(f"  {p}")


def docx_to_text(src, dst):
    from docx import Document

    doc = Document(str(src))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    dst.write_text("\n".join(parts), encoding="utf-8")


def docx_to_json(src, dst):
    from docx import Document

    doc = Document(str(src))
    result = {
        "paragraphs": [{"text": p.text, "style": p.style.name} for p in doc.paragraphs if p.text.strip()],
        "tables": [],
    }
    for ti, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        result["tables"].append({"index": ti, "rows": rows})
    dst.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")


def xlsx_to_csv(src, dst, sheet=None):
    from openpyxl import load_workbook

    wb = load_workbook(str(src), read_only=True, data_only=True)
    sname = sheet or wb.sheetnames[0]
    ws = wb[sname]
    with open(dst, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow([_cell_str(c) for c in row])
    wb.close()


def xlsx_to_json(src, dst, sheet=None):
    from openpyxl import load_workbook

    wb = load_workbook(str(src), read_only=True, data_only=True)
    sheets = [sheet] if sheet else wb.sheetnames
    result = {}
    for sname in sheets:
        if sname not in wb.sheetnames:
            continue
        ws = wb[sname]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            result[sname] = {"headers": [], "rows": []}
            continue
        headers = [_cell_str(c) for c in rows[0]]
        data = [dict(zip(headers, [_cell_str(c) for c in row])) for row in rows[1:]]
        result[sname] = {"headers": headers, "rows": data}
    wb.close()
    dst.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")


def csv_to_json(src, dst):
    with open(src, newline="", encoding="utf-8-sig") as f:
        try:
            dialect = csv.Sniffer().sniff(f.read(8192))
        except csv.Error:
            dialect = csv.excel
        f.seek(0)
        reader = csv.DictReader(f, dialect=dialect)
        rows = list(reader)
    result = {"headers": reader.fieldnames, "rows": rows, "row_count": len(rows)}
    dst.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")


def csv_to_xlsx(src, dst):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    with open(src, newline="", encoding="utf-8-sig") as f:
        try:
            dialect = csv.Sniffer().sniff(f.read(8192))
        except csv.Error:
            dialect = csv.excel
        f.seek(0)
        reader = csv.reader(f, dialect=dialect)
        for row in reader:
            ws.append(row)
    wb.save(str(dst))


def json_to_csv(src, dst):
    with open(src, encoding="utf-8") as f:
        data = json.load(f)

    # Handle both {"rows": [...]} and bare [...]
    if isinstance(data, dict):
        rows = data.get("rows", [])
    elif isinstance(data, list):
        rows = data
    else:
        print(f"Error: JSON root must be array or object with 'rows' key", file=sys.stderr)
        sys.exit(1)

    if not rows:
        dst.write_text("", encoding="utf-8")
        return

    headers = None
    if all(isinstance(row, dict) for row in rows):
        headers = []
        seen = set()
        for row in rows:
            for key in row.keys():
                key = str(key)
                if key in seen:
                    continue
                seen.add(key)
                headers.append(key)
    with open(dst, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if headers:
            writer.writerow(headers)
            for row in rows:
                writer.writerow([row.get(h, "") for h in headers])
        else:
            for row in rows:
                writer.writerow(row)


def json_to_xlsx(src, dst):
    from openpyxl import Workbook

    with open(src, encoding="utf-8") as f:
        data = json.load(f)

    sheets = _json_to_sheet_rows(data)
    wb = Workbook()
    default = wb.active
    wb.remove(default)

    for sheet_name, rows in sheets:
        ws = wb.create_sheet(_safe_sheet_name(sheet_name, wb.sheetnames))
        headers, normalized_rows = _normalize_json_rows(rows)
        if headers:
            ws.append(headers)
            for row in normalized_rows:
                ws.append([_cell_json_value(row.get(header, "")) for header in headers])
        else:
            for row in normalized_rows:
                ws.append([_cell_json_value(value) for value in row])

    dst.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(dst))


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _resolve_pages(pages_str, total):
    if not pages_str:
        return range(total)
    indices = set()
    for part in pages_str.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            indices.update(range(int(a) - 1, min(int(b), total)))
        else:
            idx = int(part) - 1
            if 0 <= idx < total:
                indices.add(idx)
    return sorted(indices)


def _cell_str(v):
    if v is None:
        return ""
    return str(v)


def _json_to_sheet_rows(data):
    if isinstance(data, dict) and "rows" not in data:
        sheet_items = []
        for name, value in data.items():
            if isinstance(value, dict) and "rows" in value:
                sheet_items.append((name, value.get("rows", [])))
            elif isinstance(value, list):
                sheet_items.append((name, value))
        if sheet_items:
            return sheet_items

    if isinstance(data, dict):
        return [("Sheet1", data.get("rows", [data]))]
    if isinstance(data, list):
        return [("Sheet1", data)]
    return [("Sheet1", [{"value": data}])]


def _normalize_json_rows(rows):
    if not rows:
        return [], []

    if all(isinstance(row, dict) for row in rows):
        headers = []
        seen = set()
        for row in rows:
            for key in row.keys():
                key = str(key)
                if key in seen:
                    continue
                seen.add(key)
                headers.append(key)
        return headers, [{str(key): value for key, value in row.items()} for row in rows]

    normalized = []
    for row in rows:
        if isinstance(row, (list, tuple)):
            normalized.append(list(row))
        else:
            normalized.append([row])
    return [], normalized


def _cell_json_value(value):
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return value
    return json.dumps(value, ensure_ascii=False)


def _safe_sheet_name(name, existing):
    cleaned = "".join("_" if ch in r'[]:*?/\\' else ch for ch in str(name or "Sheet"))
    cleaned = cleaned.strip() or "Sheet"
    base = cleaned[:31]
    candidate = base
    counter = 2
    while candidate in existing:
        suffix = f"_{counter}"
        candidate = f"{base[:31 - len(suffix)]}{suffix}"
        counter += 1
    return candidate


# ─── Routing ──────────────────────────────────────────────────────────────────

# Key: (src_ext, dst_ext) -> converter function
CONVERTERS = {
    (".pdf", ".txt"): pdf_to_text,
    (".pdf", ".json"): pdf_to_json,
    (".docx", ".txt"): docx_to_text,
    (".docx", ".json"): docx_to_json,
    (".xlsx", ".csv"): xlsx_to_csv,
    (".xlsx", ".json"): xlsx_to_json,
    (".csv", ".json"): csv_to_json,
    (".csv", ".xlsx"): csv_to_xlsx,
    (".json", ".csv"): json_to_csv,
    (".json", ".xlsx"): json_to_xlsx,
}


def main():
    parser = argparse.ArgumentParser(description="File format converter")
    parser.add_argument("input", help="Input file path")
    parser.add_argument("output", help="Output file path (use a directory for PDF->images)")
    parser.add_argument("--pages", help="Page range for PDFs (e.g. 1-5,8)")
    parser.add_argument("--sheet", help="Sheet name for XLSX source")
    parser.add_argument("--dpi", type=int, default=150, help="DPI for PDF->image (default: 150)")
    args = parser.parse_args()

    src = Path(args.input)
    dst = Path(args.output)

    if not src.exists():
        print(f"Error: input file not found: {src}", file=sys.stderr)
        sys.exit(1)

    src_ext = src.suffix.lower()
    dst_ext = dst.suffix.lower()

    # Special case: PDF -> images (output is a directory)
    if src_ext == ".pdf" and dst_ext in (".png", ""):
        out_dir = dst if dst_ext == "" else dst.parent
        pdf_to_images(src, out_dir, pages=args.pages, dpi=args.dpi)
        return

    key = (src_ext, dst_ext)
    if key not in CONVERTERS:
        print(f"Error: unsupported conversion: {src_ext} -> {dst_ext}", file=sys.stderr)
        print(f"Supported:", file=sys.stderr)
        for (s, d) in sorted(CONVERTERS.keys()):
            print(f"  {s} -> {d}", file=sys.stderr)
        sys.exit(1)

    try:
        converter = CONVERTERS[key]
        # Pass extra args for converters that accept them
        if src_ext == ".pdf":
            converter(src, dst, pages=args.pages)
        elif src_ext == ".xlsx":
            converter(src, dst, sheet=args.sheet)
        else:
            converter(src, dst)
        print(f"Converted: {src} -> {dst}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
