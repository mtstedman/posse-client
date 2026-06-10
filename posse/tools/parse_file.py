#!/usr/bin/env python3
"""Universal file parser — extracts text/data from common file formats.

Usage:
    python tools/parse_file.py <input_file> [--format text|json|csv] [--pages 1-5] [--sheet SheetName]

Supported formats:
    PDF   → text extraction (pdfplumber), table extraction
    DOCX  → paragraph text + tables
    XLSX  → sheet data as JSON or CSV
    CSV   → parsed rows as JSON
    JSON  → pretty-printed / validated
    TXT   → passthrough

Output goes to stdout. Errors go to stderr.
"""

import argparse
import csv
import io
import json
import sys
from pathlib import Path


def parse_pdf(path, pages=None, output_format="text"):
    import pdfplumber

    result = {"text": [], "tables": []}
    with pdfplumber.open(path) as pdf:
        page_range = _resolve_pages(pages, len(pdf.pages))
        for i in page_range:
            page = pdf.pages[i]
            text = page.extract_text()
            if text:
                result["text"].append({"page": i + 1, "content": text})
            tables = page.extract_tables()
            for ti, table in enumerate(tables):
                result["tables"].append({"page": i + 1, "table_index": ti, "rows": table})

    if output_format == "json":
        return json.dumps(result, indent=2, ensure_ascii=False)
    else:
        parts = []
        for entry in result["text"]:
            parts.append(f"--- Page {entry['page']} ---\n{entry['content']}")
        if result["tables"]:
            parts.append("\n--- Tables ---")
            for t in result["tables"]:
                parts.append(f"\nPage {t['page']}, Table {t['table_index']}:")
                for row in t["rows"]:
                    parts.append("\t".join(str(c or "") for c in row))
        return "\n\n".join(parts)


def parse_docx(path, output_format="text"):
    from docx import Document

    doc = Document(path)
    result = {"paragraphs": [], "tables": []}

    for para in doc.paragraphs:
        if para.text.strip():
            result["paragraphs"].append({
                "text": para.text,
                "style": para.style.name if para.style else None,
            })

    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        result["tables"].append({"table_index": ti, "rows": rows})

    if output_format == "json":
        return json.dumps(result, indent=2, ensure_ascii=False)
    else:
        parts = [p["text"] for p in result["paragraphs"]]
        if result["tables"]:
            parts.append("\n--- Tables ---")
            for t in result["tables"]:
                parts.append(f"\nTable {t['table_index']}:")
                for row in t["rows"]:
                    parts.append("\t".join(row))
        return "\n".join(parts)


def parse_xlsx(path, sheet=None, output_format="json"):
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    sheets = [sheet] if sheet else wb.sheetnames
    result = {}

    for sname in sheets:
        if sname not in wb.sheetnames:
            print(f"Warning: sheet '{sname}' not found, skipping", file=sys.stderr)
            continue
        ws = wb[sname]
        rows = []
        headers = None
        for ri, row in enumerate(ws.iter_rows(values_only=True)):
            vals = [_cell_value(c) for c in row]
            if ri == 0:
                headers = vals
            else:
                rows.append(dict(zip(headers, vals)) if headers else vals)
        result[sname] = {"headers": headers, "rows": rows, "row_count": len(rows)}

    wb.close()

    if output_format == "csv":
        out = io.StringIO()
        for sname, data in result.items():
            writer = csv.writer(out)
            if data["headers"]:
                writer.writerow(data["headers"])
            for row in data["rows"]:
                writer.writerow(row.values() if isinstance(row, dict) else row)
        return out.getvalue()
    return json.dumps(result, indent=2, ensure_ascii=False, default=str)


def parse_csv_file(path, output_format="json"):
    if output_format != "json":
        with open(path, newline="", encoding="utf-8-sig") as f:
            return f.read()

    with open(path, newline="", encoding="utf-8-sig") as f:
        # Sniff dialect
        sample = f.read(8192)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            dialect = csv.excel
        reader = csv.DictReader(f, dialect=dialect)
        rows = list(reader)

    if output_format == "json":
        return json.dumps({"headers": reader.fieldnames, "rows": rows, "row_count": len(rows)}, indent=2, ensure_ascii=False)
    return ""


def parse_json_file(path):
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)


def parse_text_file(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _resolve_pages(pages_str, total):
    if not pages_str:
        return range(total)
    parts = pages_str.split(",")
    indices = set()
    for part in parts:
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            indices.update(range(int(a) - 1, min(int(b), total)))
        else:
            idx = int(part) - 1
            if 0 <= idx < total:
                indices.add(idx)
    return sorted(indices)


def _cell_value(v):
    if v is None:
        return None
    if isinstance(v, (int, float, bool)):
        return v
    return str(v)


# ─── Main ─────────────────────────────────────────────────────────────────────

PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".csv": parse_csv_file,
    ".tsv": parse_csv_file,
    ".json": parse_json_file,
    ".txt": parse_text_file,
    ".md": parse_text_file,
    ".log": parse_text_file,
}


def main():
    parser = argparse.ArgumentParser(description="Universal file parser")
    parser.add_argument("input", help="Path to input file")
    parser.add_argument("--format", choices=["text", "json", "csv"], default=None,
                        help="Output format (default: auto based on file type)")
    parser.add_argument("--pages", help="Page range for PDFs (e.g. 1-5,8)")
    parser.add_argument("--sheet", help="Sheet name for XLSX files")
    args = parser.parse_args()

    path = Path(args.input)
    if not path.exists():
        print(f"Error: file not found: {path}", file=sys.stderr)
        sys.exit(1)

    ext = path.suffix.lower()
    if ext not in PARSERS:
        print(f"Error: unsupported file type: {ext}", file=sys.stderr)
        print(f"Supported: {', '.join(sorted(PARSERS.keys()))}", file=sys.stderr)
        sys.exit(1)

    fmt = args.format or ("json" if ext in (".xlsx", ".csv", ".tsv") else "text")

    try:
        if ext == ".pdf":
            output = parse_pdf(path, pages=args.pages, output_format=fmt)
        elif ext == ".xlsx":
            output = parse_xlsx(path, sheet=args.sheet, output_format=fmt)
        elif ext == ".docx":
            output = parse_docx(path, output_format=fmt)
        elif ext in (".csv", ".tsv"):
            output = parse_csv_file(path, output_format=fmt)
        elif ext == ".json":
            output = parse_json_file(path)
        else:
            output = parse_text_file(path)

        print(output)
    except Exception as e:
        print(f"Error parsing {path}: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
